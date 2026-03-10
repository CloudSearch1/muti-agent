"""
文档管理模块

提供文档上传、解析、分块等功能。
"""

import os
import uuid
from datetime import datetime
from typing import Any, BinaryIO

import structlog

from .chunking import ChunkerFactory, TextChunker
from .exceptions import (
    ChunkingError,
    DocumentNotFoundError,
    DocumentProcessingError,
    FileParseError,
    UnsupportedDocumentTypeError,
)
from .types import Chunk, Document, DocumentStatus, DocumentType, ChunkStrategy

logger = structlog.get_logger(__name__)

# 默认配置
DEFAULT_STORAGE_DIR = "./data/knowledge/documents"
DEFAULT_CHUNK_SIZE = 500
DEFAULT_CHUNK_OVERLAP = 50


class DocumentManager:
    """
    文档管理器

    提供文档的上传、解析、分块和管理功能。

    Example:
        >>> dm = DocumentManager()
        >>> doc = await dm.upload_document("report.pdf")
        >>> chunks = dm.get_document_chunks(doc.id)
    """

    def __init__(
        self,
        storage_dir: str = DEFAULT_STORAGE_DIR,
        chunk_strategy: ChunkStrategy = ChunkStrategy.SEMANTIC,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
    ) -> None:
        """
        初始化文档管理器

        Args:
            storage_dir: 文档存储目录
            chunk_strategy: 分块策略
            chunk_size: 分块大小
            chunk_overlap: 分块重叠
        """
        self.storage_dir = storage_dir
        self.chunk_strategy = chunk_strategy
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # 文档存储（内存中，实际应用应使用数据库）
        self._documents: dict[str, Document] = {}

        # 分块器
        self._chunker = ChunkerFactory.create(
            strategy=chunk_strategy,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

        self.logger = logger.bind(component="document_manager")

        # 确保存储目录存在
        os.makedirs(storage_dir, exist_ok=True)

    async def upload_document(
        self,
        file: BinaryIO | str,
        filename: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        """
        上传文档

        Args:
            file: 文件对象或文件路径
            filename: 文件名
            metadata: 元数据

        Returns:
            文档对象

        Raises:
            DocumentProcessingError: 处理失败
        """
        try:
            # 创建文档对象
            doc_id = str(uuid.uuid4())

            # 获取文件信息
            if isinstance(file, str):
                # 文件路径
                file_path = file
                filename = filename or os.path.basename(file_path)
                file_size = os.path.getsize(file_path)
            else:
                # 文件对象
                file_path = os.path.join(self.storage_dir, f"{doc_id}_{filename or 'unknown'}")
                content = file.read()
                with open(file_path, "wb") as f:
                    f.write(content)
                file_size = len(content)

            # 检测文档类型
            doc_type = self._detect_document_type(filename)

            # 创建文档
            document = Document(
                id=doc_id,
                title=filename or "Untitled",
                content="",
                doc_type=doc_type,
                status=DocumentStatus.PENDING,
                metadata=metadata or {},
                file_path=file_path,
                file_size=file_size,
                created_at=datetime.now(),
                updated_at=datetime.now(),
            )

            # 保存到存储
            self._documents[doc_id] = document

            self.logger.info(
                "Document uploaded",
                document_id=doc_id,
                filename=filename,
                doc_type=doc_type.value,
                file_size=file_size,
            )

            return document

        except Exception as e:
            self.logger.error("Failed to upload document", error=str(e))
            raise DocumentProcessingError(
                f"Failed to upload document: {e}",
                stage="upload",
            ) from e

    async def process_document(self, document_id: str) -> Document:
        """
        处理文档（解析和分块）

        Args:
            document_id: 文档 ID

        Returns:
            处理后的文档

        Raises:
            DocumentNotFoundError: 文档未找到
            DocumentProcessingError: 处理失败
        """
        document = self._documents.get(document_id)
        if not document:
            raise DocumentNotFoundError(document_id)

        try:
            # 更新状态
            document.status = DocumentStatus.PROCESSING
            document.updated_at = datetime.now()

            # 解析文档
            content = await self.parse_document(document.file_path, document.doc_type)
            document.content = content

            # 分块
            chunks = await self.chunk_document(
                content=content,
                document_id=document_id,
                metadata=document.metadata,
            )
            document.chunks = chunks

            # 更新状态
            document.status = DocumentStatus.READY
            document.updated_at = datetime.now()

            self.logger.info(
                "Document processed",
                document_id=document_id,
                chunks_count=len(chunks),
                content_length=len(content),
            )

            return document

        except (FileParseError, ChunkingError):
            document.status = DocumentStatus.FAILED
            document.updated_at = datetime.now()
            raise
        except Exception as e:
            document.status = DocumentStatus.FAILED
            document.updated_at = datetime.now()
            self.logger.error("Failed to process document", error=str(e))
            raise DocumentProcessingError(
                f"Failed to process document: {e}",
                document_id=document_id,
            ) from e

    async def parse_document(
        self,
        file_path: str,
        doc_type: DocumentType,
    ) -> str:
        """
        解析文档内容

        Args:
            file_path: 文件路径
            doc_type: 文档类型

        Returns:
            文档文本内容

        Raises:
            FileParseError: 解析失败
            UnsupportedDocumentTypeError: 不支持的文档类型
        """
        try:
            if doc_type == DocumentType.PDF:
                return await self._parse_pdf(file_path)
            elif doc_type == DocumentType.WORD:
                return await self._parse_word(file_path)
            elif doc_type == DocumentType.MARKDOWN:
                return await self._parse_markdown(file_path)
            elif doc_type == DocumentType.TXT:
                return await self._parse_txt(file_path)
            elif doc_type == DocumentType.HTML:
                return await self._parse_html(file_path)
            else:
                raise UnsupportedDocumentTypeError(doc_type.value)

        except (FileParseError, UnsupportedDocumentTypeError):
            raise
        except Exception as e:
            raise FileParseError(
                f"Failed to parse document: {e}",
                file_path=file_path,
                file_type=doc_type.value,
            ) from e

    async def _parse_pdf(self, file_path: str) -> str:
        """解析 PDF 文件"""
        try:
            import pypdf
        except ImportError:
            # 回退到简单文本读取
            self.logger.warning("pypdf not installed, returning empty content")
            return f"[PDF file: {os.path.basename(file_path)} - pypdf not installed]"

        try:
            text_parts = []
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            return "\n\n".join(text_parts)

        except Exception as e:
            raise FileParseError(
                f"Failed to parse PDF: {e}",
                file_path=file_path,
                file_type="pdf",
            ) from e

    async def _parse_word(self, file_path: str) -> str:
        """解析 Word 文件"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            self.logger.warning("python-docx not installed, returning empty content")
            return f"[Word file: {os.path.basename(file_path)} - python-docx not installed]"

        try:
            doc = DocxDocument(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            return "\n\n".join(text_parts)

        except Exception as e:
            raise FileParseError(
                f"Failed to parse Word: {e}",
                file_path=file_path,
                file_type="word",
            ) from e

    async def _parse_markdown(self, file_path: str) -> str:
        """解析 Markdown 文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            raise FileParseError(
                f"Failed to parse Markdown: {e}",
                file_path=file_path,
                file_type="markdown",
            ) from e

    async def _parse_txt(self, file_path: str) -> str:
        """解析纯文本文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    return f.read()
            except Exception as e:
                raise FileParseError(
                    f"Failed to parse text file: {e}",
                    file_path=file_path,
                    file_type="txt",
                ) from e
        except Exception as e:
            raise FileParseError(
                f"Failed to parse text file: {e}",
                file_path=file_path,
                file_type="txt",
            ) from e

    async def _parse_html(self, file_path: str) -> str:
        """解析 HTML 文件"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            self.logger.warning("beautifulsoup4 not installed, returning raw content")
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")

            # 移除脚本和样式
            for element in soup(["script", "style", "nav", "footer"]):
                element.decompose()

            # 获取文本
            text = soup.get_text(separator="\n", strip=True)
            return text

        except Exception as e:
            raise FileParseError(
                f"Failed to parse HTML: {e}",
                file_path=file_path,
                file_type="html",
            ) from e

    async def chunk_document(
        self,
        content: str,
        document_id: str,
        metadata: dict[str, Any] | None = None,
    ) -> list[Chunk]:
        """
        对文档内容进行分块

        Args:
            content: 文档内容
            document_id: 文档 ID
            metadata: 元数据

        Returns:
            分块列表

        Raises:
            ChunkingError: 分块失败
        """
        try:
            return self._chunker.chunk(
                text=content,
                document_id=document_id,
                metadata=metadata,
            )
        except Exception as e:
            raise ChunkingError(
                f"Failed to chunk document: {e}",
                document_id=document_id,
                strategy=self.chunk_strategy.value,
            ) from e

    def get_document(self, document_id: str) -> Document | None:
        """
        获取文档

        Args:
            document_id: 文档 ID

        Returns:
            文档对象，不存在则返回 None
        """
        return self._documents.get(document_id)

    def get_document_chunks(self, document_id: str) -> list[Chunk]:
        """
        获取文档分块

        Args:
            document_id: 文档 ID

        Returns:
            分块列表

        Raises:
            DocumentNotFoundError: 文档未找到
        """
        document = self._documents.get(document_id)
        if not document:
            raise DocumentNotFoundError(document_id)
        return document.chunks

    def list_documents(
        self,
        status: DocumentStatus | None = None,
        doc_type: DocumentType | None = None,
        limit: int = 100,
        offset: int = 0,
    ) -> list[Document]:
        """
        列出文档

        Args:
            status: 状态过滤
            doc_type: 类型过滤
            limit: 返回数量限制
            offset: 偏移量

        Returns:
            文档列表
        """
        documents = list(self._documents.values())

        # 过滤
        if status:
            documents = [d for d in documents if d.status == status]
        if doc_type:
            documents = [d for d in documents if d.doc_type == doc_type]

        # 排序（按创建时间降序）
        documents.sort(key=lambda d: d.created_at or datetime.min, reverse=True)

        # 分页
        return documents[offset : offset + limit]

    async def delete_document(self, document_id: str) -> bool:
        """
        删除文档

        Args:
            document_id: 文档 ID

        Returns:
            是否成功

        Raises:
            DocumentNotFoundError: 文档未找到
        """
        document = self._documents.get(document_id)
        if not document:
            raise DocumentNotFoundError(document_id)

        # 删除文件
        if document.file_path and os.path.exists(document.file_path):
            try:
                os.remove(document.file_path)
            except Exception as e:
                self.logger.warning(
                    "Failed to delete document file",
                    file_path=document.file_path,
                    error=str(e),
                )

        # 从存储中删除
        del self._documents[document_id]

        self.logger.info("Document deleted", document_id=document_id)
        return True

    def _detect_document_type(self, filename: str) -> DocumentType:
        """
        检测文档类型

        Args:
            filename: 文件名

        Returns:
            文档类型
        """
        ext = os.path.splitext(filename)[1].lower()

        type_map = {
            ".pdf": DocumentType.PDF,
            ".doc": DocumentType.WORD,
            ".docx": DocumentType.WORD,
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
            ".txt": DocumentType.TXT,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
        }

        return type_map.get(ext, DocumentType.UNKNOWN)

    def get_stats(self) -> dict[str, Any]:
        """
        获取统计信息

        Returns:
            统计信息字典
        """
        documents = list(self._documents.values())

        by_type: dict[str, int] = {}
        by_status: dict[str, int] = {}
        total_chunks = 0

        for doc in documents:
            # 按类型统计
            type_key = doc.doc_type.value
            by_type[type_key] = by_type.get(type_key, 0) + 1

            # 按状态统计
            status_key = doc.status.value
            by_status[status_key] = by_status.get(status_key, 0) + 1

            # 总分块数
            total_chunks += len(doc.chunks)

        return {
            "total_documents": len(documents),
            "total_chunks": total_chunks,
            "by_type": by_type,
            "by_status": by_status,
        }